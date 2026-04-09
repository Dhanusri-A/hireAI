from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor as DocxRGB, Pt
import tempfile
import os
import traceback

# ─── Helpers ──────────────────────────────────────────────────────────────────

def _sv(d, key):
    v = d.get(key, "")
    return v.strip() if isinstance(v, str) else ""

def format_list_comma(items):
    return ", ".join(str(x).strip() for x in items if str(x).strip()) or "—"

def format_education_lines(edu_list):
    lines = []
    for edu in edu_list:
        degree = edu.get("degree","").strip()
        inst   = edu.get("institution","").strip()
        year   = edu.get("year","").strip()
        spec   = edu.get("specialization","").strip()
        parts = []
        if degree: parts.append(degree)
        if inst:   parts.append(inst)
        if year:   parts.append(f"({year})")
        line = " – ".join(parts)
        if spec:   line += f" | {spec}"
        if line:   lines.append(line)
    return lines

def summarise_responsibilities(resp_list):
    """Return first 3 items joined as one paragraph sentence."""
    items = [str(r).strip() for r in resp_list if str(r).strip()][:5]
    return " ".join(items)

# ─── XML paragraph builders ────────────────────────────────────────────────────

def _new_para(style='ListParagraph'):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    ps  = OxmlElement('w:pStyle')
    ps.set(qn('w:val'), style)
    pPr.append(ps)
    p.append(pPr)
    return p

def _new_run(text, bold=False, color_hex=None, size_pt=None):
    r   = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if bold:
        rPr.append(OxmlElement('w:b'))
    if color_hex:
        clr = OxmlElement('w:color')
        clr.set(qn('w:val'), color_hex)
        rPr.append(clr)
    if size_pt:
        for tag in ('w:sz', 'w:szCs'):
            el = OxmlElement(tag)
            el.set(qn('w:val'), str(int(size_pt * 2)))
            rPr.append(el)
    if len(rPr):
        r.append(rPr)
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    return r

def _spacer():
    return OxmlElement('w:p')

# ─── Plain text replacement (for NAME, TITLE, SUMMARY, personal fields) ────────

def _replace_in_para(para, target, replacement):
    if target not in para.text:
        return False
    new_text = para.text.replace(target, replacement)
    run = next((r for r in para.runs if r.text.strip()), None)
    fmt = {}
    if run:
        fmt = {'bold': run.bold, 'italic': run.italic,
               'size': run.font.size, 'name': run.font.name}
        try: fmt['color'] = run.font.color.rgb
        except: pass
    para.clear()
    nr = para.add_run(new_text)
    nr.bold = fmt.get('bold'); nr.italic = fmt.get('italic')
    if fmt.get('size'):  nr.font.size = fmt['size']
    if fmt.get('color'): nr.font.color.rgb = fmt['color']
    if fmt.get('name'):  nr.font.name = fmt['name']
    return True

def replace_in_doc(doc, target, replacement):
    n = 0
    for p in doc.paragraphs:
        if _replace_in_para(p, target, replacement): n += 1
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if _replace_in_para(p, target, replacement): n += 1
    return n

def remove_para_containing(doc, text):
    for p in doc.paragraphs:
        if text in p.text:
            p._element.getparent().remove(p._element)
            return True
    return False

# ─── Skills: replace SKILLS placeholder with 4 labelled bullet lists ───────────
def inject_skills(doc, skills_dict):
    CATS = [
        ('project_management', 'Project Management'),
        ('technical',          'Technical'),
        ('devops',             'DevOps'),
        ('domains',            'Domains'),
    ]

    target = None
    for p in doc.paragraphs:
        if 'SKILLS' in p.text:
            target = p
            break

    if target is None:
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if 'SKILLS' in p.text:
                            _inject_skills_in_cell(cell, skills_dict, CATS)
                            return
        return

    parent = target._element.getparent()
    idx = list(parent).index(target._element)
    parent.remove(target._element)

    at = idx

    for key, label in CATS:
        items = [i.strip() for i in skills_dict.get(key, []) if i.strip()]
        if not items:
            continue

        joined = ", ".join(items)

        # Only ONE bullet paragraph per category
        ph = _new_para('ListParagraph')
        ph.append(
            _new_run(
                f"{label}: {joined}",
                bold=False,
                size_pt=10
            )
        )
        parent.insert(at, ph)
        at += 1

    parent.insert(at, _spacer())

def _inject_skills_in_cell(cell, skills_dict, CATS):
    # Remove placeholder
    for p in list(cell.paragraphs):
        if 'SKILLS' in p.text:
            p._element.getparent().remove(p._element)

    tc = cell._tc

    for key, label in CATS:
        items = [i.strip() for i in skills_dict.get(key, []) if i.strip()]
        if not items:
            continue

        joined = ", ".join(items)

        # Only ONE bullet paragraph
        ph = _new_para('ListParagraph')
        ph.append(
            _new_run(
                f"{label}: ",
                bold=True,
                color_hex="73223B",
                size_pt=10
            )
        )

        # ⚫ Normal skills
        ph.append(
            _new_run(
                joined,
                size_pt=10
            )
        )
        tc.append(ph)

    tc.append(_spacer())

# ─── Experience: bold maroon header + summary paragraph ────────────────────────

def inject_experience(doc, exp_list):
    target = None
    for p in doc.paragraphs:
        if 'EXPERIENCE' in p.text:
            target = p; break
    if not target: return
    parent = target._element.getparent()
    idx    = list(parent).index(target._element)
    parent.remove(target._element)
    at = idx
    for exp in exp_list:
        role     = exp.get("role","").strip()
        company  = exp.get("company","").strip()
        duration = exp.get("duration","").strip()
        header   = role
        if company:  header += f" | {company}"
        if duration: header += f" ({duration})"
        # Bold maroon header
        ph = _new_para('ListParagraph')
        ph.append(_new_run(header, bold=True, color_hex="73223B", size_pt=11))
        parent.insert(at, ph); at += 1
        # Summary paragraph (first 3 responsibilities as prose)
        resp   = exp.get("responsibilities", [])
        summary = summarise_responsibilities(resp)
        if summary:
            ps = _new_para('Normal')
            ps.append(_new_run(summary, size_pt=10))
            parent.insert(at, ps); at += 1
        parent.insert(at, _spacer()); at += 1

# ─── Education: bullet points ──────────────────────────────────────────────────

def inject_education(doc, edu_list):
    target = None
    for p in doc.paragraphs:
        if 'EDUCATION' in p.text:
            target = p; break
    if not target: return
    parent = target._element.getparent()
    idx    = list(parent).index(target._element)
    parent.remove(target._element)
    at = idx
    for line in format_education_lines(edu_list):
        pi = _new_para('ListParagraph')
        pi.append(_new_run(f" {line}", size_pt=10))
        parent.insert(at, pi); at += 1

# ─── Certifications: bullet points ────────────────────────────────────────────

def inject_certifications(doc, cert_list):
    target = None
    for p in doc.paragraphs:
        if 'CERTIFICATIONS' in p.text:
            target = p; break
    if not target: return
    parent = target._element.getparent()
    idx    = list(parent).index(target._element)
    parent.remove(target._element)
    at = idx
    for cert in cert_list:
        cert = str(cert).strip()
        if not cert: continue
        pi = _new_para('ListParagraph')
        pi.append(_new_run(f" {cert}", size_pt=10))
        parent.insert(at, pi); at += 1

# ─── Main generator ───────────────────────────────────────────────────────────

def generate_synpulse_docx(data: dict):
    BASE_DIR      = os.path.dirname(os.path.abspath(__file__))
    template_path = os.path.join(BASE_DIR, "format-docx.docx")
    print(f"\n=== STARTING DOCX GENERATION ===")
    if not os.path.isfile(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")
    try:
        doc      = Document(template_path)
        personal = data.get("personalInfo", {})

        dob      = _sv(personal, "dob")
        gender   = _sv(personal, "gender")
        nat      = _sv(personal, "nationality")
        marital  = _sv(personal, "marital_status")
        passport = _sv(personal, "passport")
        loc      = _sv(personal, "current_location") or _sv(personal, "location")

        # Simple text replacements
        simple = {
            "NAME":             _sv(personal, "name")  or "—",
            "TITLE":            _sv(personal, "title") or "—",
            "SUMMARY":          _sv(personal, "summary") or "—",
            "DOB":              dob     or "—",
            "GENDER":           gender  or "—",
            "NATIONALITY":      nat     or "—",
            "MARITAL_STATUS":   marital or "—",
            "PASSPORT":         passport or "—",
            "CURRENT_LOCATION": loc     or "—",
            "LANGUAGES":        format_list_comma(data.get("languages", [])),
            "HOBBIES":          format_list_comma(data.get("hobbies", [])),
        }
        for ph, val in simple.items():
            replace_in_doc(doc, ph, val)

        # Remove empty personal detail rows
        if not gender:   remove_para_containing(doc, "Gender\t\t-\t")
        if not nat:      remove_para_containing(doc, "Nationality\t\t-\t")
        if not passport: remove_para_containing(doc, "Passport\t\t-\t")
        if not dob:      remove_para_containing(doc, "Date of Birth\t-\t")
        if not marital:  remove_para_containing(doc, "Marital Status\t-\t")
        if not loc:      remove_para_containing(doc, "Current Location\t-\t")
        # If ALL personal details empty, remove section heading too
        if not any([dob, gender, nat, marital, passport, loc]):
            remove_para_containing(doc, "Personal Details")

        # Structured injections
        inject_skills(doc, data.get("skills", {}))
        inject_experience(doc, data.get("experience", []))
        inject_education(doc, data.get("education", []))
        inject_certifications(doc, data.get("certifications", []))

        out = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(out.name)
        print(f"✅ DOCX saved → {out.name}")
        return out.name
    except Exception:
        print("\n❌ FAILED"); traceback.print_exc(); raise


