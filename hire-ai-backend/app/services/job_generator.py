import os
import requests
import tempfile
import uuid
import json
import docx
from typing import Optional

from fastapi import HTTPException
from docx import Document
from fpdf import FPDF

from app.schemas.job import JobDescriptionCreate

class JobDescriptionGenerator:
    def __init__(self):
        self.api_key = os.getenv("GROQ_API_KEY")

        if not self.api_key:
            try:
                with open("groq_key.txt", "r") as f:
                    self.api_key = f.read().strip()
            except FileNotFoundError:
                self.api_key = None

        self.api_url = "https://api.groq.com/openai/v1/chat/completions"
    
    def generate_job_description(
        self,
        data: JobDescriptionCreate,
    ) -> dict:
        if not self.api_key:
            raise HTTPException(
                status_code=500,
                detail="Groq API key not configured",
            )

        prompt = self._build_prompt(data)

        payload = {
            "model": "llama-3.1-8b-instant",
            "messages": [
                {
                    "role": "system",
                    "content": (
                        "You are a professional HR content writer. "
                        "Generate clear, structured, and realistic job descriptions "
                        "with headings and bullet points."
                    ),
                },
                {"role": "user", "content": prompt},
            ],
            "temperature": 0.7,
            "max_tokens": 1500,
        }

        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json",
        }

        try:
            response = requests.post(
                self.api_url,
                headers=headers,
                json=payload,
                timeout=30,
            )

            if response.status_code != 200:
                raise HTTPException(
                    status_code=response.status_code,
                    detail=response.text,
                )

            content = response.json()["choices"][0]["message"]["content"]
            try:
                return json.loads(content)
            except json.JSONDecodeError:
                raise HTTPException(
                    status_code=500,
                    detail="LLM returned invalid JSON",
                )

        except requests.RequestException as exc:
            raise HTTPException(
                status_code=500,
                detail=f"Groq API request failed: {exc}",
            )


    def _build_prompt(self, data: JobDescriptionCreate) -> str:
        return f"""
You are a senior HR professional who writes structured job descriptions for global companies.

IMPORTANT CORE RULE
All generated content MUST depend strictly on the following inputs:
- Experience Level
- Tone Style
- Skills
- Responsibilities (if provided)

If a level is "Entry Level", the job description MUST NOT contain senior-level expectations.

If a level is "Senior" or "Lead", the job description MUST contain leadership, architecture, and mentoring responsibilities.

Do NOT contradict the experience level.

------------------------------------------------

STRICT OUTPUT RULES

- Return ONLY valid JSON
- Do NOT include markdown
- Do NOT include explanations
- JSON must start with {{ and end with }}

OUTPUT FORMAT

{{
  "about_the_role": "string (120-200 words)",

  "minimum_experience": "string (EXPLAIN CLEARLY using years of experience and expectations based on level)",

  "key_responsibilities": ["string"],
  "required_skills": ["string"],
  "preferred_skills": ["string"],
  "education_requirements": ["string"],
  "qualities_and_requirements": ["string"],
  "tools_and_technologies": ["string"],

  "work_environment": "string (60-120 words)",

  "benefits_and_perks": ["string"],

  "about_us": "string (80-150 words)"
}}

------------------------------------------------

ARRAY RULES

- key_responsibilities: 5-10 items
- required_skills: 5-10 items
- preferred_skills: 5-7 items
- education_requirements: 5-7 items
- qualities_and_requirements: 5-10 items
- tools_and_technologies: 5-7 items
- benefits_and_perks: 5-7 items

------------------------------------------------

EXPERIENCE LEVEL LOGIC (MANDATORY)

Entry Level
0-1 years of experience
Focus on learning, assisting team members, and executing assigned tasks.

Junior Level
1-3 years of experience
Can independently develop smaller modules and collaborate with the team.

Mid Level
3-6 years of experience
Own features, contribute to architecture discussions, mentor juniors occasionally.

Senior Level
6-10 years of experience
Lead projects, mentor engineers, review code, and drive technical decisions.

Lead / Principal
10+ years of experience
Provide technical leadership, guide engineering teams, manage stakeholders, and define system architecture.

IMPORTANT RULE
The "minimum_experience" field MUST contain a detailed sentence describing the required experience using years.
Example style:
"Candidates should typically have 3–6 years of professional experience in backend development, demonstrating the ability to design scalable APIs and contribute to system architecture."

Never output text like:
"Mid Level experience required"

------------------------------------------------

RESPONSIBILITY ALIGNMENT RULE

Responsibilities MUST match the level:

Entry
Assist, learn, support senior engineers, complete assigned tasks.

Junior
Develop modules, debug issues, collaborate with team members.

Mid
Design features, mentor juniors, optimize performance.

Senior
Lead technical initiatives, review code, guide engineers.

Lead
Define architecture, lead teams, manage stakeholders.

If the input contains responsibilities, expand them but keep the same level alignment.

------------------------------------------------

SKILL ALIGNMENT RULE

All skill-related sections MUST align with:
- Skills provided in input
- Responsibilities
- Experience level

Example:
Entry Level → basic frameworks, learning mindset
Senior Level → architecture, scalability, leadership

------------------------------------------------

TONE APPLICATION RULE

Tone must influence EVERY TEXT SECTION including:
- about_the_role
- minimum_experience
- responsibilities
- requirements
- work_environment
- benefits

Tone definitions:

Professional
Formal corporate tone used on LinkedIn and job boards.

Engaged
Friendly, motivating tone encouraging applicants.

Casual
Relaxed conversational style but still clear and structured.

Inclusive
Welcoming language encouraging diverse candidates and avoiding bias.

Dynamic
Energetic startup-style tone highlighting innovation and growth.

IMPORTANT
Tone changes WRITING STYLE only.
Section titles must NEVER change.

------------------------------------------------

ABOUT US RULE

The "about_us" section MUST be written strictly based on the provided About Company information.
Do NOT invent unrelated company descriptions.

------------------------------------------------

JOB INPUT

Job Title: {data.job_title}
Company Name: {data.company_name}
Department: {data.department or "Not specified"}
Location: {data.location or "Not specified"}
Experience Level: {data.level or "Not specified"}
Tone Style: {data.tone_style or "Professional"}
Skills: {data.skills or "Generate relevant skills based on job title"}
Responsibilities: {data.responsibilities or "Generate responsibilities based on role"}
About Company: {data.about_company or "Generate company overview"}
Additional Context: {data.additional_data or "None"}
Description: {data.input_description or "None"}

------------------------------------------------

FINAL GENERATION RULE

Every section MUST be consistent with:
- the experience level
- the provided skills
- the responsibilities
- the selected tone style

The job description must look like a realistic posting from a professional company on LinkedIn or a major job board.
"""

    def save_to_word(self, text: str, filename: Optional[str] = None) -> str:
        """Save job description to Word document and return file path"""
        if not filename:
            filename = f"Job_Description_{uuid.uuid4().hex[:8]}.docx"
        
        doc = Document()
        doc.add_heading("Job Description", level=1)
        doc.add_paragraph(text)
        
        # Save to temporary file
        temp_dir = tempfile.gettempdir()
        file_path = os.path.join(temp_dir, filename)
        doc.save(file_path)
        
        return file_path
    
    def save_to_pdf(self, text: str, filename: Optional[str] = None) -> str:
        """Save job description to PDF and return file path"""
        if not filename:
            filename = f"Job_Description_{uuid.uuid4().hex[:8]}.pdf"
        
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        
        for line in text.split('\n'):
            if line.strip():  # Skip empty lines
                pdf.multi_cell(0, 10, line)
        
        # Save to temporary file
        temp_dir = tempfile.gettempdir()
        file_path = os.path.join(temp_dir, filename)
        pdf.output(file_path)
        
        return file_path

job_generator = JobDescriptionGenerator()

